"""
Extract readable text from common file types found in SharePoint:
  Excel (.xlsx / .xls), CSV, Word (.docx), PDF, plain text / JSON / Markdown.

Returns a plain-text string that can be sent to Claude.
"""

import io
import os

# Maximum characters to send to Claude per file to avoid hitting context limits.
# Roughly ~50 K tokens for a 200 K char excerpt.
MAX_CHARS = 200_000


def extract_text(file_content: bytes, filename: str) -> str:
    """Dispatch to the right extractor based on file extension."""
    ext = os.path.splitext(filename)[1].lower()

    extractors = {
        ".xlsx": _excel,
        ".xls": _excel,
        ".csv": _csv,
        ".docx": _word,
        ".doc": _word,
        ".pdf": _pdf,
        ".txt": _plain,
        ".md": _plain,
        ".json": _plain,
        ".xml": _plain,
        ".html": _plain,
        ".htm": _plain,
    }

    extractor = extractors.get(ext)
    if extractor is None:
        raise ValueError(
            f"Unsupported file type '{ext}'. "
            f"Supported types: {', '.join(sorted(extractors))}"
        )

    text = extractor(file_content)
    return _truncate(text, filename)


# ─── Extractors ───────────────────────────────────────────────────────────────

def _excel(content: bytes) -> str:
    import openpyxl

    wb = openpyxl.load_workbook(io.BytesIO(content), data_only=True, read_only=True)
    parts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            if any(cell is not None for cell in row):
                rows.append("\t".join("" if cell is None else str(cell) for cell in row))
        if rows:
            parts.append(f"=== Sheet: {sheet_name} ===\n" + "\n".join(rows))
    wb.close()
    return "\n\n".join(parts) if parts else "(empty workbook)"


def _csv(content: bytes) -> str:
    import pandas as pd

    for enc in ("utf-8", "latin-1", "cp1252"):
        try:
            df = pd.read_csv(io.BytesIO(content), encoding=enc)
            return df.to_string(index=False)
        except UnicodeDecodeError:
            continue
    raise ValueError("Could not decode CSV — unrecognised encoding.")


def _word(content: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(content))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    table_parts = []
    for tbl in doc.tables:
        rows = ["\t".join(cell.text for cell in row.cells) for row in tbl.rows]
        table_parts.append("\n".join(rows))

    text = "\n".join(paragraphs)
    if table_parts:
        text += "\n\n=== Tables ===\n" + "\n\n".join(table_parts)
    return text or "(empty document)"


def _pdf(content: bytes) -> str:
    import pdfplumber

    pages = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for i, page in enumerate(pdf.pages, 1):
            page_text = page.extract_text()
            if page_text and page_text.strip():
                pages.append(f"--- Page {i} ---\n{page_text}")
    return "\n\n".join(pages) if pages else "(no extractable text in PDF)"


def _plain(content: bytes) -> str:
    for enc in ("utf-8", "latin-1", "cp1252"):
        try:
            return content.decode(enc)
        except UnicodeDecodeError:
            continue
    raise ValueError("Could not decode text file — unrecognised encoding.")


# ─── Helpers ──────────────────────────────────────────────────────────────────

def _truncate(text: str, filename: str) -> str:
    if len(text) <= MAX_CHARS:
        return text
    truncated = text[:MAX_CHARS]
    omitted = len(text) - MAX_CHARS
    notice = (
        f"\n\n[⚠️  File '{filename}' was truncated. "
        f"{omitted:,} characters ({omitted // 4:,} estimated tokens) were omitted. "
        f"Ask about specific sections if you need the rest.]"
    )
    return truncated + notice
